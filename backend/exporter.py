"""Word (.docx) export: the refined 3-column claim chart plus a refinement
change log appendix — the artifact an analyst would file or share with counsel."""
from __future__ import annotations

import io
import re

from backend.chart_store import ChartStore
from backend.parsers import clean_text

_STRENGTH_LABEL = {"strong": "Strong", "moderate": "Moderate", "weak": "Weak"}


def export_filename(store: ChartStore) -> str:
    """Safe .docx filename derived from the chart title."""
    stem = re.sub(r"[^A-Za-z0-9]+", "_", store.current.title).strip("_")[:40]
    return f"{stem or 'claim_chart'}_refined.docx"


def export_docx(store: ChartStore) -> bytes:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt, RGBColor

    chart = store.current
    doc = Document()

    title = doc.add_heading(chart.title, level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    doc.add_paragraph(
        "Claim chart refined with iLumos AI-assisted analysis. All AI-suggested "
        "changes below were reviewed and approved by the analyst.")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Patent Claim Element", "Accused Product Feature (Evidence)",
               "AI Reasoning", "Evidence Strength")
    for cell, text in zip(table.rows[0].cells, headers):
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(10)

    for row in chart.rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (row.element, row.feature, row.reasoning,
                                      _STRENGTH_LABEL.get(row.strength, row.strength))):
            para = cell.paragraphs[0]
            # clean_text: python-docx rejects XML-invalid control characters,
            # and chart text may arrive via chat paths that bypass the parsers
            run = para.add_run(clean_text(text))
            run.font.size = Pt(9)

    log = store.change_log()
    if log:
        doc.add_heading("Appendix — Refinement Change Log", level=2)
        for i, entry in enumerate(log, start=1):
            doc.add_paragraph(f"{i}. {entry}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
