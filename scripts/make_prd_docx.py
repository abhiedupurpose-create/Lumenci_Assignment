"""Render docs/2_PRD.md into docs/PRD.docx (1-page-style Word document) so the
PRD can be uploaded/shared per the assignment's Word/Google-doc requirement.

Run:  python scripts/make_prd_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "2_PRD.md"
DST = ROOT / "docs" / "PRD.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
_SPAN_RE = re.compile(r"\*\*(?P<bold>.+?)\*\*|\*(?P<ital>[^*\n]+?)\*")


def add_runs(paragraph, text: str) -> None:
    """Write text with **bold** and *italic* spans as proper runs."""
    pos = 0
    for m in _SPAN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("bold") is not None:
            paragraph.add_run(m.group("bold")).bold = True
        else:
            paragraph.add_run(m.group("ital")).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def main() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # One-page discipline: tight margins, compact type, minimal heading air.
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.0  # template default ~1.08 costs a page
    for list_style in ("List Bullet", "List Number"):
        s = doc.styles[list_style]
        s.font.size = Pt(9)
        s.paragraph_format.space_after = Pt(1)
        s.paragraph_format.line_spacing = 1.0

    def _tighten(heading, size):
        for run in heading.runs:
            run.font.color.rgb = NAVY
            run.font.size = Pt(size)
        heading.paragraph_format.space_before = Pt(5)
        heading.paragraph_format.space_after = Pt(2)

    for line in lines:
        line = line.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            _tighten(doc.add_heading(line[2:], level=0), 15)
        elif line.startswith("## "):
            _tighten(doc.add_heading(line[3:], level=1), 11.5)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_runs(p, line)

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    sys.exit(main())
